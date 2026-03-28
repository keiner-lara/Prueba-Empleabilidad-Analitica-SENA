import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Please install python-docx")
    sys.exit(1)

AZUL_OSCURO = RGBColor(0x1A, 0x23, 0x7E)
AZUL_MEDIO  = RGBColor(0x15, 0x65, 0xC0)
GRIS_TEXTO  = RGBColor(0x37, 0x47, 0x4F)
BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc: Document, text: str, level: int = 1, color: RGBColor = None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color or AZUL_OSCURO
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        else:
            run.font.size = Pt(12)
    return p

def body(doc: Document, text: str, size: int = 11, color: RGBColor = None, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or GRIS_TEXTO
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p

def insert_image(doc: Document, abs_path: str, caption: str, width=Inches(5.0)):
    if not os.path.exists(abs_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(abs_path, width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

def portada(doc: Document, titulo: str, subtitulo: str):
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"SENA — PROGRAMA DE FORMACIÓN")
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL_MEDIO
    run.bold = True
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(titulo)
    run3.font.size = Pt(26)
    run3.font.color.rgb = AZUL_OSCURO
    run3.bold = True
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(subtitulo)
    run4.font.size = Pt(13)
    run4.font.color.rgb = GRIS_TEXTO
    run4.italic = True
    for _ in range(4): doc.add_paragraph()
    
    datos = [
        ("Aprendiz:",       "Keiner Lara"),
        ("Programa SENA:",  "Desarrollo de software"),
        ("Ruta Riwi:",      "Analítica de Datos"),
        ("Proyecto:",       "Plataforma de Postulación a Empleabilidad"),
        ("Fecha:",          "Marzo 2026"),
    ]
    t = doc.add_table(rows=len(datos), cols=2)
    t.style = 'Table Grid'
    for i, (label, val) in enumerate(datos):
        c0, c1 = t.rows[i].cells
        set_cell_bg(c0, '1A237E')
        set_cell_bg(c1, 'F5F5F5')
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.color.rgb = BLANCO
        r0.bold = True
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = GRIS_TEXTO
    doc.add_page_break()

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

portada(
    doc,
    titulo="Solución de software funcional",
    subtitulo="Evidencia 3 - Norma 220501096"
)

heading(doc, "1. Introducción", level=1)
body(doc, "El presente documento formaliza la entrega de la solución de software desarrollada bajo los estándares de la norma de competencia 'Desarrollar la solución de software de acuerdo con el diseño y metodologías de desarrollo'.")

heading(doc, "2. Acceso al Código Fuente", level=1)
body(doc, "El código completo, incluyendo la lógica de ingeniería de datos (ETL), scripts de base de datos y el portal frontend, se encuentra alojado en el repositorio oficial:")
body(doc, "• Repositorio GitHub: https://github.com/keiner-lara/Prueba-Empleabilidad-Analitica-SENA.git", bold=True, color=AZUL_MEDIO)
body(doc, "• Rama principal: main (Código estable y listo para producción).")

heading(doc, "3. Aplicación en Vivo (Despliegue)", level=1)
body(doc, "Para validar el funcionamiento del diseño propuesto y la navegación entre pantallas, la aplicación ha sido desplegada en la infraestructura de Vercel:")
body(doc, "• URL de la App: https://prueba-empleabilidad-analitica-sena-ayrc61pon.vercel.app/", bold=True, color=AZUL_MEDIO)
body(doc, "• Tecnologías de Despliegue: CI/CD automatizado sincronizado con la rama main.")

heading(doc, "4. Componentes de la Solución", level=1)
body(doc, "La solución integra los siguientes productos entregados previamente:")
body(doc, "• Base de Datos: Script SQL robusto con modelo Star Schema (bbdd.sql).")
body(doc, "• Ingeniería de Datos: Scripts Python integrales para limpieza y conexión (scripts/).")
body(doc, "• Documentación Técnica: Manuales de usuario y de código fuente detallados (normas/).")

heading(doc, "5. Conclusiones y Cumplimiento", level=1)
body(doc, "La solución cumple satisfactoriamente con todos los requisitos funcionales y no funcionales establecidos por el Programa SENA, demostrando competencias avanzadas en el desarrollo de software y análisis de datos en la nube.")

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'normas', 'Solución de software.docx')
out_path = os.path.abspath(out_path)
doc.save(out_path)
print(f"Documento generado exitosamente en: {out_path}")
