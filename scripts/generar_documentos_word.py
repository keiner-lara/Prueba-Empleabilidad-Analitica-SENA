"""
Genera los dos documentos Word para la prueba SENA:
  - NORMA_220501095_Diseno_KL.docx  (NORMA 95 — Diseño)
  - NORMA_220501096_Desarrollo_KL.docx  (NORMA 96 — Desarrollo)

Requisito: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── HELPERS ────────────────────────────────────────────────────────────────

AZUL_OSCURO = RGBColor(0x1A, 0x23, 0x7E)
AZUL_MEDIO  = RGBColor(0x15, 0x65, 0xC0)
GRIS_TEXTO  = RGBColor(0x37, 0x47, 0x4F)
BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)
FONDO_TABLA = RGBColor(0xE3, 0xF2, 0xFD)

BASE = os.path.dirname(os.path.abspath(__file__))


def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def heading(doc: Document, text: str, level: int = 1, color: RGBColor = None):
    """Agrega un heading con color personalizado."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color or AZUL_OSCURO
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        else:
            run.font.size = Pt(12)
    return p


def body(doc: Document, text: str, size: int = 11, color: RGBColor = None, bold=False, italic=False):
    """Agrega un párrafo de texto normal."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or GRIS_TEXTO
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def code_block(doc: Document, code: str):
    """Agrega un bloque de código con fuente monoespaciada y fondo gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    # Fondo del párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F8E9')
    pPr.append(shd)
    return p


def insert_image(doc: Document, rel_path: str, caption: str, width=Inches(5.5)):
    """Inserta una imagen con pie de foto centrado."""
    full_path = os.path.join(BASE, rel_path)
    if not os.path.exists(full_path):
        body(doc, f"[IMAGEN NO ENCONTRADA: {rel_path}]", italic=True, color=RGBColor(0xB0, 0x00, 0x20))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(full_path, width=width)
    # Pie de foto
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)


def divider(doc: Document):
    """Línea horizontal separadora."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1565C0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(10)


def portada(doc: Document, norma: str, titulo: str, subtitulo: str):
    """Genera la portada del documento."""
    # Espacio superior
    for _ in range(3):
        doc.add_paragraph()

    # Norma badge
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"SENA — PROGRAMA DE FORMACIÓN")
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL_MEDIO
    run.bold = True

    # Número de norma
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(norma)
    run2.font.size = Pt(13)
    run2.font.color.rgb = GRIS_TEXTO
    run2.bold = True

    doc.add_paragraph()

    # Título principal
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(titulo)
    run3.font.size = Pt(26)
    run3.font.color.rgb = AZUL_OSCURO
    run3.bold = True

    # Subtítulo
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(subtitulo)
    run4.font.size = Pt(13)
    run4.font.color.rgb = GRIS_TEXTO
    run4.italic = True

    for _ in range(4):
        doc.add_paragraph()

    # Datos del autor
    datos = [
        ("Aprendiz:",       "Keiner Lara"),
        ("Programa SENA:",  "Tecnología en Análisis y Desarrollo de Sistemas de Información"),
        ("Ruta Riwi:",      "Analítica de Datos"),
        ("Proyecto:",       "Plataforma de Postulación a Empleabilidad"),
        ("Fecha:",          "Marzo 2026"),
    ]
    t = doc.add_table(rows=len(datos), cols=2)
    t.style = 'Table Grid'
    for i, (label, val) in enumerate(datos):
        c0, c1 = t.rows[i].cells
        set_cell_bg(c0, '1A237E')
        set_cell_bg(c1, 'F5F5F5')
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.color.rgb = BLANCO
        r0.bold = True
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = GRIS_TEXTO

    doc.add_page_break()


def tabla_simple(doc: Document, headers: list, rows: list):
    """Crea una tabla estilizada con encabezado azul."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, '1565C0')
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Filas
    for ri, row in enumerate(rows):
        bg = 'E3F2FD' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = GRIS_TEXTO
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# NORMA 220501095 — DISEÑO
# ═══════════════════════════════════════════════════════════════════════════════

def generar_norma_95():
    doc = Document()
    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    portada(
        doc,
        norma="NORMA DE COMPETENCIA 220501095",
        titulo="Diseño de Solución de Software",
        subtitulo="Plataforma Analítica de Ventas LATAM",
    )

    # ── EVIDENCIA 1: DOCUMENTO DE DISEÑO DE SOFTWARE ──────────────────────────
    heading(doc, "EVIDENCIA 1 — Documento de Diseño de Software", level=1)
    divider(doc)

    heading(doc, "1.1 Introducción", level=2)
    body(doc, (
        "Este documento describe el diseño de la solución analítica para la plataforma de empleabilidad, "
        "desarrollada sobre un dataset de ventas LATAM con 2,000,000 de registros. El objetivo es presentar "
        "el enfoque completo de análisis de datos, desde el proceso ETL hasta la visualización interactiva "
        "y el portal web de resultados."
    ))

    heading(doc, "1.2 Problema a Resolver", level=2)
    body(doc, (
        "La organización requiere transformar un dataset de ventas con inconsistencias —errores ortográficos "
        "en categorías, valores nulos en columnas críticas, tipos de datos mixtos— en información confiable "
        "para la toma de decisiones estratégicas. El proceso debe garantizar calidad de datos, modelado "
        "analítico y visualización clara para stakeholders de negocio."
    ))

    heading(doc, "1.3 Objetivos del Sistema", level=2)
    for obj in [
        "Limpiar y normalizar el dataset de 2,000,000 filas garantizando integridad referencial.",
        "Construir un modelo de datos estrella en PostgreSQL para consultas analíticas eficientes.",
        "Generar insights accionables sobre ventas, segmentos de clientes y distribución regional.",
        "Presentar resultados en dashboards de Power BI y en un portal web accesible sin instalaciones.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(obj)
        run.font.size = Pt(11)
        run.font.color.rgb = GRIS_TEXTO

    doc.add_paragraph()

    heading(doc, "1.4 Actores del Sistema", level=2)
    tabla_simple(doc,
        headers=["Actor", "Rol", "Interacción con el Sistema"],
        rows=[
            ["Analista de Datos",          "Usuario técnico",      "Ejecuta ETL, valida calidad, carga a BD"],
            ["Stakeholder de Negocio",      "Usuario consumidor",   "Consulta dashboards y portal web"],
            ["DBA / DevOps",               "Administrador",        "Configura PostgreSQL y variables de entorno"],
        ]
    )

    heading(doc, "1.5 Requisitos Funcionales y No Funcionales", level=2)
    tabla_simple(doc,
        headers=["Tipo", "Requisito", "Criterio de Aceptación"],
        rows=[
            ["Funcional",     "Limpieza automática de datos",              "0 nulos en columnas críticas"],
            ["Funcional",     "Generación de dataset limpio",              "Datos_limpio.csv con 1,951,391 filas"],
            ["Funcional",     "Carga a modelo estrella",                   "4 tablas pobladas sin duplicados"],
            ["Funcional",     "Consultas SQL para KPIs",                   "Top 5, avg ticket, participación"],
            ["Funcional",     "Portal web con KPIs e insights",            "Carga en navegador sin servidor"],
            ["No funcional",  "Reproducible en ambiente local",            "Documentado con requirements.txt"],
            ["No funcional",  "Rendimiento aceptable para 2M filas",       "ETL completa en < 5 minutos"],
            ["No funcional",  "Código documentado y mantenible",           "Comentarios en cada bloque lógico"],
        ]
    )

    heading(doc, "1.6 Arquitectura del Sistema", level=2)
    body(doc, "La solución sigue una arquitectura de capas lineales (pipeline analítico):")
    for paso, desc in [
        ("Ingesta",        "Datos.csv — 2,000,000 filas de ventas LATAM"),
        ("ETL",            "limpieza.py — Python + Pandas (correcciones, imputación, tipo_producto)"),
        ("Almacenamiento", "PostgreSQL — modelo estrella (4 dimensiones + 1 tabla hechos)"),
        ("Visualización",  "Power BI Desktop — dashboards ejecutivos con 4 vistas"),
        ("Presentación",   "portal/index.html — portal HTML/CSS con Chart.js"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run0 = p.add_run(f"{paso}: ")
        run0.bold = True
        run0.font.size = Pt(11)
        run0.font.color.rgb = AZUL_MEDIO
        run1 = p.add_run(desc)
        run1.font.size = Pt(11)
        run1.font.color.rgb = GRIS_TEXTO

    doc.add_paragraph()

    heading(doc, "1.7 Módulos del Sistema", level=2)
    tabla_simple(doc,
        headers=["Módulo", "Archivo", "Responsabilidad"],
        rows=[
            ["ETL / Limpieza",      "limpieza.py",              "15 pasos de transformación, exporta CSV limpio"],
            ["Carga de Datos",      "conexion_carga_datos.py",  "Inserta en 4 tablas del modelo estrella"],
            ["Esquema BD",          "bbdd.sql",                 "DDL de tablas, constraints e índices"],
            ["Visualización BI",    "Dashboards/*.png",         "4 capturas de Power BI con KPIs"],
            ["Portal Web",          "portal/index.html",        "KPIs, Top 5 con Chart.js, insights, galería"],
        ]
    )

    heading(doc, "1.8 Tecnologías Utilizadas", level=2)
    tabla_simple(doc,
        headers=["Tecnología", "Versión", "Uso en el Proyecto"],
        rows=[
            ["Python",      "3.x",    "Motor ETL, scripting de carga"],
            ["Pandas",      "≥1.5",   "Transformaciones sobre 2M filas"],
            ["PostgreSQL",  "13+",    "Data Warehouse con modelo estrella"],
            ["SQLAlchemy",  "≥2.0",   "ORM y conexión con transacciones"],
            ["Power BI",    "Desktop","Dashboards interactivos ejecutivos"],
            ["HTML/CSS",    "5 / 3",  "Portal web de resultados"],
            ["Bootstrap",   "5.3",    "Sistema de grids y componentes UI"],
            ["Chart.js",    "4.4",    "Gráfico de barras interactivo"],
        ]
    )

    heading(doc, "1.9 Justificación Técnica", level=2)
    body(doc, (
        "La arquitectura híbrida propuesta —Python como motor ETL, PostgreSQL como repositorio analítico, "
        "Power BI como capa de visualización empresarial y HTML/CSS como portal de publicación— fue "
        "seleccionada por su equilibrio entre rendimiento, reproducibilidad y adopción en la industria. "
        "Python con Pandas permite procesar datasets de más de 2 millones de filas con pipelines explícitos "
        "y trazables, donde cada transformación queda documentada como un paso numerado en el código fuente."
    ))
    body(doc, (
        "El modelo de datos estrella implementado en PostgreSQL garantiza que las consultas analíticas sean "
        "eficientes incluso a gran escala, siguiendo el estándar Kimball para Data Warehousing. La separación "
        "en tablas de dimensión y una tabla de hechos central facilita la agregación por múltiples ejes sin "
        "redundancia, y es extensible sin modificar el esquema de hechos."
    ))
    body(doc, (
        "La decisión de exponer los resultados mediante un portal web estático responde a la necesidad de "
        "democratizar el acceso a los insights sin requerir herramientas especializadas. El portal complementa "
        "el dashboard de Power BI ofreciendo los KPIs, rankings y hallazgos estratégicos en un formato "
        "universalmente accesible desde cualquier navegador."
    ))

    heading(doc, "1.10 Métricas de Calidad de Datos (Antes vs. Después)", level=2)
    tabla_simple(doc,
        headers=["Métrica", "Dataset Original", "Dataset Limpio"],
        rows=[
            ["Total de filas",                   "2,000,000",     "1,951,391"],
            ["Filas eliminadas",                  "—",             "48,609  (~2.4%)"],
            ["Columnas con nulos críticos",        "4",             "0"],
            ["Errores ortográficos en categorías","15+ variantes", "0"],
            ["Tipos de datos incorrectos",         "3 columnas",    "0"],
            ["Columnas derivadas añadidas",        "0",             "4 (year, month, year_month, segmento_cliente)"],
        ]
    )

    doc.add_page_break()

    # ── EVIDENCIA 2: DIAGRAMAS UML ─────────────────────────────────────────────
    heading(doc, "EVIDENCIA 2 — Diagramas UML", level=1)
    divider(doc)

    heading(doc, "2.1 Diagrama de Casos de Uso", level=2)
    body(doc, (
        "El diagrama de casos de uso describe los actores del sistema (Analista de Datos y Stakeholder de "
        "Negocio) y las interacciones que cada uno tiene con el sistema analítico."
    ))
    insert_image(doc, "docs/diagrams/uml_casos_uso.png",
                 "Figura 1 – Diagrama de Casos de Uso", width=Inches(5))

    heading(doc, "2.2 Diagrama de Componentes", level=2)
    body(doc, (
        "El diagrama de componentes muestra la estructura modular del sistema, incluyendo las dependencias "
        "entre el ETL (Python), la base de datos (PostgreSQL), Power BI y el portal web."
    ))
    insert_image(doc, "docs/diagrams/uml_componentes.png",
                 "Figura 2 – Diagrama de Componentes", width=Inches(5))

    heading(doc, "2.3 Diagrama de Actividad", level=2)
    body(doc, (
        "El diagrama de actividad ilustra el flujo de ejecución del pipeline ETL: desde la carga del CSV "
        "original, pasando por cada transformación, hasta la exportación del dataset limpio y la carga a BD."
    ))
    insert_image(doc, "docs/diagrams/uml_actividad.png",
                 "Figura 3 – Diagrama de Actividad (Pipeline ETL)", width=Inches(5))

    doc.add_page_break()

    # ── EVIDENCIA 3: PROTOTIPO / WIREFRAMES ───────────────────────────────────
    heading(doc, "EVIDENCIA 3 — Prototipo (Wireframes)", level=1)
    divider(doc)

    body(doc, (
        "Los wireframes a continuación representan las tres pantallas principales del portal analítico, "
        "diseñadas siguiendo los principios de UX de claridad, jerarquía visual y accesibilidad de datos "
        "para usuarios no técnicos (stakeholders de negocio)."
    ))

    heading(doc, "3.1 Pantalla 1 — Inicio / KPIs", level=2)
    body(doc, (
        "Vista principal con 4 tarjetas de métricas clave: Total Sales ($33.83B), Avg Ticket ($17.34K), "
        "Units Sold (10.73M) y LATAM Hubs (3). Incluye barra de navegación a las demás secciones."
    ))
    insert_image(doc, "docs/wireframes/wireframe_inicio.png",
                 "Figura 4 – Wireframe Pantalla 1: KPIs Principales", width=Inches(5.5))

    heading(doc, "3.2 Pantalla 2 — Top 5 Productos", level=2)
    body(doc, (
        "Combina una tabla de datos con el ranking de los 5 productos con mayores ventas y un gráfico de "
        "barras horizontal interactivo (Chart.js) para facilitar la comparación visual."
    ))
    insert_image(doc, "docs/wireframes/wireframe_top5.png",
                 "Figura 5 – Wireframe Pantalla 2: Top 5 Productos", width=Inches(5.5))

    heading(doc, "3.3 Pantalla 3 — Insights & Galería de Dashboards", level=2)
    body(doc, (
        "Presenta 3 hallazgos estratégicos en tarjetas estructuradas (Evidencia / Impacto / Recomendación) "
        "y una galería de las 4 capturas del dashboard de Power BI con lightbox interactivo."
    ))
    insert_image(doc, "docs/wireframes/wireframe_insights.png",
                 "Figura 6 – Wireframe Pantalla 3: Insights y Galería", width=Inches(5.5))

    doc.add_page_break()

    # ── EVIDENCIA 4: MODELO DE BASE DE DATOS ──────────────────────────────────
    heading(doc, "EVIDENCIA 4 — Modelo de Base de Datos (ER)", level=1)
    divider(doc)

    heading(doc, "4.1 Descripción del Modelo Estrella", level=2)
    body(doc, (
        "El modelo de datos sigue el patrón Kimball (esquema estrella), donde una tabla de hechos central "
        "(fact_ventas) referencia cuatro tablas de dimensión: dim_fecha, dim_producto, dim_region y "
        "dim_cliente. Este diseño optimiza las consultas OLAP al minimizar los JOINs necesarios para "
        "calcular KPIs agregados."
    ))

    heading(doc, "4.2 Tablas del Modelo", level=2)
    tabla_simple(doc,
        headers=["Tabla", "Tipo", "Columnas Clave", "Descripción"],
        rows=[
            ["dim_fecha",    "Dimensión", "id_fecha, fecha, year, month, year_month", "Calendario analítico"],
            ["dim_producto", "Dimensión", "id_producto, producto, tipo_producto",     "Catálogo de SKUs"],
            ["dim_region",   "Dimensión", "id_region, ciudad, pais",                  "Geografía LATAM"],
            ["dim_cliente",  "Dimensión", "id_cliente, tipo_cliente, segmento, tipo_venta, id_region", "Segmentos"],
            ["fact_ventas",  "Hechos",    "id_venta, id_fecha, id_producto, id_cliente, cantidad, precio_unitario, descuento, costo_envio, total", "Transacciones"],
        ]
    )

    heading(doc, "4.3 Diagrama Entidad-Relación", level=2)
    insert_image(doc, "docs/diagrams/modelo_er.png",
                 "Figura 7 – Diagrama ER (Modelo Estrella PostgreSQL)", width=Inches(5))

    out = os.path.join(BASE, "NORMA_220501095_Diseno_KL.docx")
    doc.save(out)
    print(f"✅ NORMA 95 generada: {out}")
    return out


# ═══════════════════════════════════════════════════════════════════════════════
# NORMA 220501096 — DESARROLLO
# ═══════════════════════════════════════════════════════════════════════════════

def generar_norma_96():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    portada(
        doc,
        norma="NORMA DE COMPETENCIA 220501096",
        titulo="Desarrollo de Solución de Software",
        subtitulo="Plataforma Analítica de Ventas LATAM",
    )

    # ── EVIDENCIA 1: DOCUMENTO TÉCNICO ────────────────────────────────────────
    heading(doc, "EVIDENCIA 1 — Documento Técnico de Código Fuente", level=1)
    divider(doc)

    heading(doc, "1.1 Estructura del Proyecto", level=2)
    code_block(doc, (
        "Prueba Empleabilidad Edison Avellaneda/\n"
        "├── Datos.csv                   # Dataset original (2,000,000 filas)\n"
        "├── Datos_limpio.csv            # Dataset limpio (1,951,391 filas)\n"
        "├── limpieza.py                 # Pipeline ETL — 15 pasos documentados\n"
        "├── conexion_carga_datos.py     # Carga a PostgreSQL (modelo estrella)\n"
        "├── bbdd.sql                    # DDL + consultas SQL analíticas\n"
        "├── requirements.txt            # Dependencias Python\n"
        "├── .env                        # DATABASE_URL (fuera del código)\n"
        "├── portal/\n"
        "│   ├── index.html              # Portal web con Chart.js\n"
        "│   └── styles.css\n"
        "├── Dashboards/                 # 4 capturas de Power BI\n"
        "└── docs/                       # Documentación técnica y wireframes"
    ))

    heading(doc, "1.2 Flujo del Pipeline ETL", level=2)
    code_block(doc,
        "Datos.csv → limpieza.py → Datos_limpio.csv → conexion_carga_datos.py → PostgreSQL\n"
        "                                                                              ↓\n"
        "                                                              Power BI  /  portal/index.html"
    )

    heading(doc, "1.3 Fragmento 1 — Limpieza de Texto y Correcciones Ortográficas", level=2)
    body(doc,
        "Fragmento extraído de limpieza.py — pasos 3 y 5:",
        bold=True, color=AZUL_MEDIO
    )
    code_block(doc, (
        "# 3. LIMPIEZA DE TEXTO\n"
        "# Elimina caracteres especiales y normaliza espacios en los campos de texto\n"
        "def limpiar_texto(valor):\n"
        "    if pd.isna(valor):\n"
        "        return valor                     # Preserva NaN para manejo posterior\n"
        "    if isinstance(valor, str):\n"
        "        valor = re.sub(r'[@#*$%^&(){}[]<>~`]', '', valor)\n"
        "        valor = re.sub(r'\\s+', ' ', valor)  # Colapsa espacios múltiples\n"
        "        return valor.strip()\n"
        "    return valor\n"
        "\n"
        "# 5. CORRECCIONES ORTOGRÁFICAS\n"
        "# Mapeo: versión incorrecta → versión canónica del producto\n"
        "correcciones_producto = {\n"
        "    'Yogur': 'Yogurt',  'Yogurth': 'Yogurt',  'Gasosa': 'Gaseosa',\n"
        "    'Mantequila': 'Mantequilla', 'Galleta': 'Galletas', 'Cafe': 'Café',\n"
        "}\n"
        "if 'producto' in df.columns:\n"
        "    df['producto'] = df['producto'].replace(correcciones_producto)"
    ))
    body(doc, (
        "Impacto: Sin este paso, 'Yogur', 'Yogurth' y 'Yogurtt' serían 3 SKUs distintos, "
        "inflando artificialmente el catálogo y distorsionando los KPIs de ventas por producto."
    ), italic=True)

    heading(doc, "1.4 Fragmento 2 — Mapeo Producto → Tipo de Producto", level=2)
    body(doc, "Fragmento extraído de limpieza.py — paso 6:", bold=True, color=AZUL_MEDIO)
    code_block(doc, (
        "# 6. NORMALIZAR TIPO_PRODUCTO SEGÚN PRODUCTO\n"
        "# Garantiza que cada producto pertenezca a exactamente una categoría\n"
        "mapeo_tipo_producto = {\n"
        "    'Leche': 'Lácteo',    'Queso': 'Lácteo',   'Mantequilla': 'Lácteo',\n"
        "    'Yogurt': 'Lácteo',   'Té': 'Bebida',       'Café': 'Bebida',\n"
        "    'Gaseosa': 'Bebida',  'Chocolate': 'Snack', 'Galletas': 'Snack',\n"
        "    'Pan': 'Abarrotes',   'Cereal': 'Abarrotes','Arepa': 'Alimento Perecedero',\n"
        "}\n"
        "\n"
        "if 'producto' in df.columns and 'tipo_producto' in df.columns:\n"
        "    # .map() aplica el diccionario; .fillna() preserva valores para productos\n"
        "    # no listados (safe fallback — no rompe el pipeline con productos nuevos)\n"
        "    df['tipo_producto'] = df['producto'].map(mapeo_tipo_producto)\\\n"
        "                                       .fillna(df['tipo_producto'])"
    ))

    heading(doc, "1.5 Fragmento 3 — Carga a Dimensiones y Hechos (PostgreSQL)", level=2)
    body(doc, "Fragmento extraído de conexion_carga_datos.py:", bold=True, color=AZUL_MEDIO)
    code_block(doc, (
        "# Transacción atómica: si cualquier INSERT falla, hace rollback automático\n"
        "with engine.begin() as conn:\n"
        "\n"
        "    # DIM_REGION — inserta pares únicos ciudad/país\n"
        "    regiones = df[['ciudad', 'pais']].drop_duplicates()\n"
        "    conn.execute(text(\"\"\"\n"
        "        INSERT INTO dim_region (ciudad, pais) VALUES (:ciudad, :pais)\n"
        "        ON CONFLICT (ciudad, pais) DO NOTHING  -- Idempotente\n"
        "    \"\"\"), regiones.to_dict(orient='records'))\n"
        "\n"
        "    # Recupera IDs generados para el JOIN con fact_ventas\n"
        "    region_map = pd.read_sql('SELECT id_region, ciudad, pais FROM dim_region', conn)\n"
        "    df = df.merge(region_map, on=['ciudad', 'pais'], how='left')\n"
        "\n"
        "    # FACT_VENTAS — carga solo filas con todas las FK presentes\n"
        "    fact = df[['id_fecha','id_producto','id_cliente',\n"
        "               'cantidad','precio_unitario','descuento','costo_envio','total']].dropna()\n"
        "    conn.execute(text(\"\"\"\n"
        "        INSERT INTO fact_ventas (id_fecha, id_producto, id_cliente,\n"
        "                                cantidad, precio_unitario, descuento, costo_envio, total)\n"
        "        VALUES (:id_fecha,:id_producto,:id_cliente,\n"
        "                :cantidad,:precio_unitario,:descuento,:costo_envio,:total)\n"
        "    \"\"\"), fact.to_dict(orient='records'))"
    ))
    body(doc, (
        "Buenas prácticas aplicadas: uso de ON CONFLICT para idempotencia, transacción atómica "
        "con engine.begin(), credenciales externalizadas en .env (nunca hardcodeadas), "
        "y dropna() como validación final antes de insertar en hechos."
    ), italic=True)

    heading(doc, "1.6 Tecnologías y Dependencias", level=2)
    tabla_simple(doc,
        headers=["Tecnología", "Rol", "Justificación"],
        rows=[
            ["Python 3.x",    "Motor ETL",              "Ecosistema de datos maduro, Pandas para 2M+ filas"],
            ["Pandas",        "Transformaciones",        "API vectorizada, alto rendimiento sin SQL"],
            ["SQLAlchemy",    "ORM / Conexión",          "Abstracción de BD, transacciones atómicas"],
            ["PostgreSQL",    "Data Warehouse",          "ACID, modelo estrella, consultas OLAP"],
            ["Power BI",      "Visualización BI",        "Dashboards interactivos para stakeholders"],
            ["HTML/CSS",      "Portal web",              "Accesible sin instalación, responsive"],
            ["Chart.js",      "Gráficos interactivos",   "Top 5 productos con animación y tooltips"],
            ["python-dotenv", "Seguridad",               "Credenciales fuera del código fuente"],
        ]
    )

    heading(doc, "1.7 Resultados Reales del Pipeline ETL", level=2)
    tabla_simple(doc,
        headers=["Métrica", "Valor"],
        rows=[
            ["Filas originales",           "2,000,000"],
            ["Filas finales limpias",       "1,951,391"],
            ["Total Sales",                "$33.83B"],
            ["Avg Ticket",                 "$17.34K"],
            ["Units Sold",                 "10.73M"],
            ["Países cubiertos",           "Colombia, Brasil, México"],
            ["Tiempo estimado ETL",        "~2–3 min en hardware estándar"],
        ]
    )

    doc.add_page_break()

    # ── EVIDENCIA 2: INSTRUCTIVO / MANUAL DE USUARIO ──────────────────────────
    heading(doc, "EVIDENCIA 2 — Instructivo de Uso (Manual de Usuario)", level=1)
    divider(doc)

    heading(doc, "2.1 Requisitos del Sistema", level=2)
    tabla_simple(doc,
        headers=["Componente", "Versión Mínima", "Notas"],
        rows=[
            ["Python",            "3.8+",   "Con pip instalado"],
            ["PostgreSQL",        "13+",    "Con usuario y base de datos creados"],
            ["Navegador web",     "Moderno","Chrome, Firefox, Edge — para el portal"],
            ["Power BI Desktop",  "Reciente","Opcional — para explorar el .pbix"],
            ["RAM",               "8 GB",   "Recomendado por el dataset de 2M filas"],
        ]
    )

    heading(doc, "2.2 Instalación y Configuración", level=2)
    body(doc, "Paso 1 — Instalar dependencias Python:", bold=True)
    code_block(doc, "pip install -r requirements.txt")

    body(doc, "Paso 2 — Configurar el archivo .env en la raíz del proyecto:", bold=True)
    code_block(doc, "DATABASE_URL=postgresql://usuario:contraseña@localhost:5432/empleabilidad")

    body(doc, "Paso 3 — Crear el esquema de base de datos:", bold=True)
    code_block(doc, "psql -U usuario -d empleabilidad -f bbdd.sql")

    heading(doc, "2.3 Ejecución del Pipeline", level=2)

    body(doc, "Paso 1 — Ejecutar limpieza de datos:", bold=True)
    code_block(doc, "python limpieza.py")

    body(doc, "Salida esperada en consola:", italic=True)
    code_block(doc, (
        "Filas originales: 2000000\n"
        "============================================================\n"
        "DATA CLEANING SUMMARY\n"
        "============================================================\n"
        "Filas finales: 1951391\n"
        "\n"
        "Valores nulos por columna:\n"
        "fecha              0\n"
        "producto           0\n"
        "cantidad           0\n"
        "precio_unitario    0\n"
        "total              0\n"
        "dtype: int64\n"
        "\n"
        "Valores únicos (columnas categóricas):\n"
        "- producto: 12 únicos\n"
        "- tipo_producto: 5 únicos\n"
        "- pais: 3 únicos\n"
        "\n"
        "OK Archivo generado: Datos_limpio.csv"
    ))

    body(doc, "Paso 2 — Cargar datos a PostgreSQL:", bold=True)
    code_block(doc, "python conexion_carga_datos.py")

    body(doc, "Salida esperada en consola:", italic=True)
    code_block(doc, (
        "Columnas detectadas en el CSV:\n"
        "Index(['fecha', 'year', 'month', 'year_month', 'producto', 'tipo_producto',\n"
        "       'cantidad', 'precio_unitario', 'descuento', 'costo_envio', 'total',\n"
        "       'tipo_venta', 'tipo_cliente', 'segmento_cliente', 'ciudad', 'pais'],\n"
        "      dtype='object')\n"
        "\n"
        "ETL COMPLETADO CON ÉXITO"
    ))

    body(doc, "Paso 3 — Abrir el portal web:", bold=True)
    body(doc, "Navegar a la carpeta portal/ y abrir index.html con doble clic (o arrastrar al navegador).")

    heading(doc, "2.4 Solución de Problemas Comunes", level=2)
    tabla_simple(doc,
        headers=["Error", "Causa Probable", "Solución"],
        rows=[
            ["ModuleNotFoundError: pandas",         "Dependencias no instaladas",           "pip install -r requirements.txt"],
            ["OperationalError: could not connect", "PostgreSQL no activo o .env incorrecto","Revisar servicio y archivo .env"],
            ["FileNotFoundError: Datos.csv",        "Ejecutar desde carpeta incorrecta",     "Ejecutar desde la raíz del proyecto"],
            ["El portal no muestra imágenes",       "Rutas relativas incorrectas",           "Abrir index.html desde su carpeta portal/"],
        ]
    )

    doc.add_page_break()

    # ── EVIDENCIA 3: SOLUCIÓN DE SOFTWARE FUNCIONAL ───────────────────────────
    heading(doc, "EVIDENCIA 3 — Solución de Software Funcional", level=1)
    divider(doc)

    body(doc, (
        "La solución de software se compone de scripts Python funcionales, un esquema SQL operativo, "
        "un portal web interactivo y dashboards de Power BI. A continuación se presenta la evidencia "
        "visual de la solución funcionando."
    ))

    heading(doc, "3.1 Dashboard Power BI — Resumen Ejecutivo", level=2)
    insert_image(doc, "Dashboards/Dashboard.png",
                 "Figura 8 – Dashboard Power BI: Resumen Ejecutivo con KPIs", width=Inches(5.5))

    heading(doc, "3.2 Dashboard Power BI — Top 5 Productos y Clientes", level=2)
    insert_image(doc, "Dashboards/Top5.png",
                 "Figura 9 – Dashboard Power BI: Ranking Top 5 Productos y Segmentos", width=Inches(5.5))

    heading(doc, "3.3 Dashboard Power BI — Tendencia Temporal", level=2)
    insert_image(doc, "Dashboards/Tendencia temporal.png",
                 "Figura 10 – Dashboard Power BI: Análisis de Tendencia Temporal", width=Inches(5.5))

    heading(doc, "3.4 Dashboard Power BI — Participación por Categoría", level=2)
    insert_image(doc, "Dashboards/Participación por categoía.png",
                 "Figura 11 – Dashboard Power BI: Market Share por Tipo de Producto", width=Inches(5.5))

    heading(doc, "3.5 Resumen de Funcionalidades Implementadas", level=2)
    tabla_simple(doc,
        headers=["Componente", "Estado", "Observaciones"],
        rows=[
            ["limpieza.py",              "✅ Funcional", "15 pasos ETL, genera Datos_limpio.csv"],
            ["conexion_carga_datos.py",  "✅ Funcional", "Carga idempotente a 4 tablas PostgreSQL"],
            ["bbdd.sql",                 "✅ Completo",  "DDL con constraints, FK e índices de performance"],
            ["portal/index.html",        "✅ Funcional", "Chart.js, lightbox, KPIs, insights"],
            ["Dashboards (×4)",          "✅ Completo",  "Exportados de Power BI en alta resolución"],
            ["requirements.txt",         "✅ Completo",  "Dependencias versionadas"],
            [".env",                     "✅ Seguro",    "Credenciales externalizadas, no en el código"],
        ]
    )

    out = os.path.join(BASE, "NORMA_220501096_Desarrollo_KL.docx")
    doc.save(out)
    print(f"✅ NORMA 96 generada: {out}")
    return out


# ── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generando documentos Word SENA...")
    generar_norma_95()
    generar_norma_96()
    print("\n🎓 Ambos documentos Word generados con éxito.")
